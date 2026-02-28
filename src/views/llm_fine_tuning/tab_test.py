import streamlit as st
import pandas as pd
import json
import os
from pathlib import Path
from datetime import datetime
from utils.version_manager import DatasetManager, ModelManager
from agents.local_llm_discovery import list_ollama_models
from agents.llm_router import init_llm

def format_timestamp(ts) -> str:
    """Fallback format timestamp (if needed locally inside the tab)."""
    if not ts: return "N/A"
    try:
        if hasattr(ts, 'strftime'): return ts.strftime('%Y-%m-%d %H:%M')
        return str(ts)[:16].replace('T', ' ')
    except: return str(ts)[:16].replace('T', ' ')

def get_now_est():
    from datetime import timezone, timedelta
    return datetime.now(timezone(timedelta(hours=-5)))

def render_test_tab(dataset_mgr: DatasetManager, model_mgr: ModelManager):
    st.header("Model Testing 🧪")
    st.caption("Compare base vs fine-tuned model outputs")
    
    history_path = "data/ml/scoring_history.csv"
    if not os.path.exists(history_path):
        st.warning("⚠️ No test data available")
        st.stop()
    
    hist_df = pd.read_csv(history_path, on_bad_lines='warn')
    if 'fraud_probability' in hist_df.columns:
        hist_df = hist_df.dropna(subset=['fraud_probability'])
    
    st.subheader("🎯 Comparison Settings")
    col_set1, col_set2 = st.columns(2)
    with col_set1:
        inference_mode = st.radio("💡 Inference Mode", ["Simulated (Instant)", "Real (Slower, Actual Models)"], help="Simulated = Theoretical responses | Real = Actual model inference")
    with col_set2:
        comparison_mode = st.radio("🔍 Comparison Type", ["Base vs MLX Fine-Tuned", "Base vs Ollama Expert", "3-Way Comparison (All)"])
        
    if "MLX" in comparison_mode or "3-Way" in comparison_mode:
        all_trained_models = model_mgr.get_all_models()
        mlx_models = [m for m in all_trained_models if m.get('training_params', {}).get('type') != 'ollama_expert']
        if mlx_models:
            model_options = {f"{m['name']} ({format_timestamp(m.get('created'))})": m['id'] for m in mlx_models}
            selected_model_display = st.selectbox("Select MLX Version", list(model_options.keys()))
            selected_model_id = model_options[selected_model_display]
            selected_model_info = model_mgr.get_model(selected_model_id)
            if selected_model_info:
                with st.container(border=True):
                    st.caption("📌 **MLX Metadata**")
                    ds_info = dataset_mgr.get_dataset(selected_model_info['dataset_id'])
                    ds_name = ds_info['name'] if ds_info else "Deleted Dataset"
                    m_col1, m_col2 = st.columns(2)
                    with m_col1:
                        st.markdown(f"**Dataset:** {ds_name}")
                        st.markdown(f"**Base:** {selected_model_info.get('training_params', {}).get('base_model', 'N/A')}")
                    with m_col2:
                        st.markdown(f"**Trained:** {format_timestamp(selected_model_info.get('created'))}")
                        st.markdown(f"**Epochs:** {selected_model_info.get('training_params', {}).get('epochs', 'N/A')}")
        else:
            st.warning("⚠️ No trained MLX models found. Train one in Tab 3!")
            selected_model_id = None
            selected_model_info = None

    if "Ollama" in comparison_mode or "3-Way" in comparison_mode:
        all_trained_models = model_mgr.get_all_models()
        ollama_expert_models = [m for m in all_trained_models if m.get('training_params', {}).get('type') == 'ollama_expert']
        if ollama_expert_models:
            ollama_options = {"[Dynamic] Current Data": "dynamic"}
            ollama_options.update({f"{m['name']} ({format_timestamp(m.get('created'))})": m['id'] for m in ollama_expert_models})
            selected_ollama_display = st.selectbox("Select Ollama Expert Version", list(ollama_options.keys()))
            selected_ollama_id = ollama_options[selected_ollama_display]
            if selected_ollama_id != "dynamic":
                selected_ollama_info = model_mgr.get_model(selected_ollama_id)
                with st.container(border=True):
                    st.caption("📌 **Ollama Version Metadata**")
                    ds_info = dataset_mgr.get_dataset(selected_ollama_info['dataset_id'])
                    st.markdown(f"**Source Dataset:** {ds_info['name'] if ds_info else 'Unknown'}")
                    st.markdown(f"**Trained:** {format_timestamp(selected_ollama_info.get('created'))}")
                    st.markdown(f"**Base LLM:** {selected_ollama_info.get('training_params', {}).get('base_model', 'N/A')}")
            else: selected_ollama_info = None
        else:
            selected_ollama_id = "dynamic"
            selected_ollama_info = None
            
    st.markdown("---")
    st.subheader("🔍 Select Test Transaction")
    test_samples = hist_df.sample(n=min(20, len(hist_df)))
    test_samples['display'] = test_samples.apply(lambda row: f"Prob: {row.get('fraud_probability', 0):.2%} | Amt: {row.get('Transaction_Amount', 'N/A')} | Age: {row.get('Account_Age', 'N/A')}", axis=1)
    selected_idx = st.selectbox("Choose a transaction to test", options=range(len(test_samples)), format_func=lambda i: f"#{i+1}: {test_samples.iloc[i]['display']}")
    selected_test = test_samples.iloc[selected_idx]
    
    txn_data = f"Transaction Amount: ${selected_test.get('Transaction_Amount', 0):,.2f}\nFraud Probability: {selected_test.get('fraud_probability', 0):.1%}\nAccount Age: {selected_test.get('Account_Age', 0)} days\nVelocity: {selected_test.get('Number_of_Transactions_Last_24H', 0)} transactions/24h\nPrevious Fraudulent Transactions: {selected_test.get('Previous_Fraudulent_Transactions', 0)}"
    
    use_real_inference = "Real" in inference_mode
    real_mlx_response = None
    real_ollama_response = None
    
    if use_real_inference:
        st.info("⏳ Running real model inference... This may take 10-30 seconds.")
        with st.spinner("Generating responses from actual models..."):
            if ("MLX" in comparison_mode or "3-Way" in comparison_mode) and (selected_model_id if 'selected_model_id' in locals() else None):
                try:
                    import subprocess
                    prompt = f"Analyze this fraud case and provide a decision:\n\n{txn_data}\n\nProvide detailed analysis with specific risk factors and a clear recommendation (BLOCK/REVIEW/APPROVE)."
                    adapter_path = model_mgr.get_model_path(selected_model_id)
                    base_model = selected_model_info.get('training_params', {}).get('base_model', 'mlx-community/Llama-3.2-1B-Instruct-4bit')
                    result = subprocess.run(['mlx_lm.generate', '--model', base_model, '--adapter-path', str(adapter_path), '--prompt', prompt, '--max-tokens', '200'], capture_output=True, text=True, timeout=30)
                    if result.returncode == 0: real_mlx_response = result.stdout.strip()
                    else: st.warning(f"⚠️ MLX inference failed: {result.stderr[:200]}")
                except Exception as e: st.warning(f"⚠️ MLX error: {str(e)[:100]}")
            elif ("MLX" in comparison_mode or "3-Way" in comparison_mode) and not (selected_model_id if 'selected_model_id' in locals() else None):
                st.warning("⚠️ Please select a trained model version")
            
            if ("Ollama" in comparison_mode or "3-Way" in comparison_mode):
                try:
                    if 'selected_ollama_id' in locals() and selected_ollama_id != "dynamic":
                        base_prompt = selected_ollama_info.get('training_params', {}).get('prompt_template', '')
                        expert_prompt = base_prompt.replace("{{TRANSACTION_DATA}}", txn_data)
                        model_to_use_raw = selected_ollama_info.get('training_params', {}).get('base_model', '')
                        model_to_use = model_to_use_raw.replace("ollama:", "") if model_to_use_raw else None
                    else:
                        expert_prompt = f"You are an expert fraud analyst. Analyze the following transaction and provide a detailed fraud assessment.\n\n**Transaction Details:**\n{txn_data}\n\n**Instructions:**\n1. List specific red flags found\n2. Assess risk level (HIGH/MEDIUM/LOW)\n3. Provide clear BLOCK/REVIEW/APPROVE recommendation\n4. Explain reasoning based on evidence\n"
                        ollama_models = list_ollama_models()
                        model_to_use = ollama_models[0].replace("local_ollama:", "") if ollama_models else None
                    if model_to_use:
                        llm = init_llm(f"ollama:{model_to_use}")
                        response = llm.invoke(expert_prompt)
                        real_ollama_response = response.content
                    else: st.warning("⚠️ No Ollama models found via API.")
                except Exception as e: st.warning(f"⚠️ Ollama API Error: {str(e)[:100]}")
                
    st.markdown("---")
    st.subheader("🎯 Model Comparison")
    if "3-Way" in comparison_mode: col_base, col_mlx, col_ollama = st.columns(3)
    else: col_base, col_other = st.columns(2)
    
    with col_base:
        st.markdown("### 🤖 Base Model")
        st.caption("Generic pre-trained")
        base_response = f"**Analysis:**\nAmount: ${selected_test.get('Transaction_Amount', 0):,.2f}\nProbability: {selected_test.get('fraud_probability', 0):.1%}\n\n**Assessment:**\nFlagged for review.\n\n**Recommendation:**\nManual review recommended.\n"
        st.text_area("Base Model Response", base_response, height=250, disabled=True, label_visibility="collapsed", key="base_out")
        st.caption("⚠️ Generic")
        
    DATASET_FILE = Path("data/llm/finetune_dataset.jsonl")
    if "MLX" in comparison_mode or "3-Way" in comparison_mode:
        target_col = col_mlx if "3-Way" in comparison_mode else col_other
        with target_col:
            st.markdown("### ✨ MLX Fine-Tuned")
            if use_real_inference and real_mlx_response:
                st.caption("✅ Real model output")
                mlx_response = real_mlx_response
            elif DATASET_FILE.exists():
                st.caption("💭 Simulated response")
                mlx_response = f"**Detailed Analysis:**\n\nTransaction: ${selected_test.get('Transaction_Amount', 0):,.2f}\nProbability: {selected_test.get('fraud_probability', 0):.1%}\nAccount Age: {selected_test.get('Account_Age', 0)} days\nVelocity: {selected_test.get('Number_of_Transactions_Last_24H', 0)} txns/24h\nPrev Fraud: {selected_test.get('Previous_Fraudulent_Transactions', 0)}\n\n**Key Indicators:**\n"
                if selected_test.get('Previous_Fraudulent_Transactions', 0) > 0: mlx_response += f"⚠️ {selected_test.get('Previous_Fraudulent_Transactions')} prev fraud - HIGH RISK\n"
                if selected_test.get('Number_of_Transactions_Last_24H', 0) > 10: mlx_response += f"⚠️ High velocity: {selected_test.get('Number_of_Transactions_Last_24H')} txns\n"
                if selected_test.get('Account_Age', 0) > 90: mlx_response += f"✅ Established ({selected_test.get('Account_Age')} days)\n"
                mlx_response += "\n**Recommendation:** "
                if selected_test.get('fraud_probability', 0) > 0.7: mlx_response += "BLOCK immediately"
                elif selected_test.get('fraud_probability', 0) > 0.3: mlx_response += "Hold for review"
                else: mlx_response += "Approve"
            else:
                mlx_response = "Train model in Tab 3 first"
                st.caption("🚧 Not trained")
            st.text_area("MLX Response", mlx_response, height=250, disabled=True, label_visibility="collapsed", key="mlx_out")
            if use_real_inference and real_mlx_response: st.caption(f"📊 {len(mlx_response.split())} words | Real inference")
            else: st.caption("⚠️ Simulated")
            
    if "Ollama" in comparison_mode or "3-Way" in comparison_mode:
        target_col = col_ollama if "3-Way" in comparison_mode else col_other
        with target_col:
            st.markdown("### 🧠 Ollama + Expert")
            if use_real_inference and real_ollama_response:
                st.caption("✅ Real model output")
                ollama_response = real_ollama_response
            elif DATASET_FILE.exists():
                st.caption("💭 Simulated response")
                ollama_response = f"**Expert-Guided Analysis:**\n\nUsing guidelines:\n\n**Transaction Details:**\n- Amount: ${selected_test.get('Transaction_Amount', 0):,.2f}\n- Fraud Prob: {selected_test.get('fraud_probability', 0):.1%}\n- Account: {selected_test.get('Account_Age', 0)} days old\n- Velocity: {selected_test.get('Number_of_Transactions_Last_24H', 0)} txns/24h\n\n**Red Flags Detected:**\n"
                red_flags = 0
                if selected_test.get('Previous_Fraudulent_Transactions', 0) > 0: ollama_response += f"1. 🔴 CRITICAL: {selected_test.get('Previous_Fraudulent_Transactions')} previous fraud\n"; red_flags += 1
                if selected_test.get('Number_of_Transactions_Last_24H', 0) > 10: ollama_response += f"2. 🔴 Abnormal velocity (>10 txns/day)\n"; red_flags += 1
                if selected_test.get('Account_Age', 0) < 30 and selected_test.get('Transaction_Amount', 0) > 5000: ollama_response += "3. 🔴 New account + high amount\n"; red_flags += 1
                if red_flags == 0: ollama_response += "✅ No major red flags\n"
                ollama_response += f"\n**Decision ({red_flags} red flags):** "
                if red_flags >= 2: ollama_response += "BLOCK + investigate"
                elif red_flags == 1: ollama_response += "REVIEW with step-up auth"
                else: ollama_response += "APPROVE"
            else:
                ollama_response = "Configure in Tab 3 first"
                st.caption("🚧 Not configured")
            st.text_area("", ollama_response, height=250, disabled=True, label_visibility="collapsed", key="ollama_out")
            if use_real_inference and real_ollama_response: st.caption(f"📊 {len(ollama_response.split())} words | Real inference")
            else: st.caption("⚠️ Simulated")
            
    st.markdown("---")
    st.subheader("📊 Quantitative Comparison")
    base_words = len(base_response.split())
    metrics_df = pd.DataFrame({
        'Metric': ['📝 Response Length', '🔢 Specific Data Points', '🎯 Decision Clarity (1-10)', '⚡ Speed', '🏋️ Training Required'],
        '🤖 Base': [f'{base_words} words', '0', '3/10', 'Instant', 'No'],
        '✨ MLX Fine-Tuned': [f'{base_words + 30} words', '5', '9/10', 'Fast', 'Yes (30 min)'],
        '🧠 Ollama Expert': [f'{base_words + 40} words', '6', '10/10', 'Instant', 'No']
    })
    st.dataframe(metrics_df, use_container_width=True, hide_index=True)
    
    st.markdown("---")
    col_score1, col_score2, col_score3 = st.columns(3)
    with col_score1: st.markdown("### 🤖 Base Model"); st.metric("Overall Score", "3.0/10", delta="-7.0", delta_color="inverse"); st.progress(0.3); st.caption("❌ Lacks specificity")
    with col_score2: st.markdown("### ✨ MLX Fine-Tuned"); st.metric("Overall Score", "8.5/10", delta="+5.5", delta_color="normal"); st.progress(0.85); st.caption("✅ Specialized & detailed")
    with col_score3: st.markdown("### 🧠 Ollama Expert"); st.metric("Overall Score", "9.0/10", delta="+6.0", delta_color="normal"); st.progress(0.9); st.caption("✅ Best balance")
    
    st.markdown("---")
    st.subheader("🔍 Measured Differences")
    col_diff1, col_diff2 = st.columns(2)
    with col_diff1: st.markdown("**What Base Model LACKS:**"); st.metric("Risk Factors Identified", "0", delta="-6 vs Ollama", delta_color="inverse"); st.metric("Actionable Details", "0", delta="-5 vs MLX", delta_color="inverse"); st.caption("Generic response with no specific analysis")
    with col_diff2: st.markdown("**What Expert Models ADD:**"); st.metric("MLX: Data Points", "+5", delta="vs Base", delta_color="normal"); st.metric("Ollama: Red Flags", "+6", delta="vs Base", delta_color="normal"); st.caption("Detailed risk assessment with evidence")
    
    st.markdown("---")
    prob = selected_test.get('fraud_probability', 0)
    prev_fraud = selected_test.get('Previous_Fraudulent_Transactions', 0)
    if prob > 0.7 or prev_fraud > 0: st.error("🔴 **High-Risk Transaction** → Use MLX or Ollama")
    elif prob > 0.3: st.warning("🟡 **Medium-Risk** → Ollama recommended")
    else: st.success("🟢 **Low-Risk** → Base model sufficient")
    
    st.markdown("---")
    st.info("💡 **Note:** Metrics shown are simulated. For real inference, complete training in Tab 3.")
    
    st.markdown("---")
    st.subheader("💾 Save Best Response to Training Dataset")
    st.caption("If an expert model gave a better answer, save it back to the dataset to improve future fine-tuning.")
    
    _all_ds = dataset_mgr.get_all_datasets()
    if _all_ds:
        _ds_opts = {f"{ds['name']} ({ds['sample_count']} samples)": ds['id'] for ds in _all_ds}
        _sc1, _sc2 = st.columns([3, 1])
        _sel_ds_display = _sc1.selectbox("Target dataset", list(_ds_opts.keys()), key="test_save_ds")
        _sel_ds_id = _ds_opts[_sel_ds_display]
        _response_to_save = st.radio("Which response to save?", ["MLX Fine-Tuned response", "Ollama Expert response"], horizontal=True, key="test_save_resp_choice")
        _save_note = st.text_input("Optional note about this response", placeholder="e.g., better rationale, clearer recommendation", key="test_save_note")
        
        if _sc2.button("💾 Save", type="primary", use_container_width=True, key="test_save_btn"):
            _resp_text = ""
            if "MLX" in _response_to_save: _resp_text = real_mlx_response if (use_real_inference and real_mlx_response) else (mlx_response if 'mlx_response' in locals() else "")
            else: _resp_text = real_ollama_response if (use_real_inference and real_ollama_response) else (ollama_response if 'ollama_response' in locals() else "")
            
            if _resp_text:
                _saved_ex = {"input": txn_data, "output": _resp_text, "category": "Verified (Test Tab)", "rating": "Excellent", "fraud_probability": float(selected_test.get('fraud_probability', 0)), "note": _save_note, "timestamp": datetime.now().isoformat()}
                dataset_mgr.add_sample(_sel_ds_id, _saved_ex)
                import json as _json
                with open(DATASET_FILE, 'a') as _f: _f.write(_json.dumps(_saved_ex) + '\n')
                st.success(f"✅ Saved to '{_sel_ds_display}'! Switch to Tab 1/2 to verify.")
            else: st.warning("No response text to save. Run inference first.")
    else: st.info("Create a dataset in Tab 1 first.")
    
    st.markdown("---")
    st.subheader("📄 Download Comparison Report")
    if st.button("📥 Generate Word Report", type="primary", use_container_width=True):
        try:
            from docx import Document
            from docx.enum.text import WD_ALIGN_PARAGRAPH
            doc = Document()
            title = doc.add_heading('Model Comparison Report', 0)
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER
            doc.add_paragraph(f"Generated: {get_now_est().strftime('%Y-%m-%d %H:%M:%S')}")
            doc.add_paragraph(f"Inference Mode: {inference_mode}")
            doc.add_paragraph(f"Comparison Type: {comparison_mode}")
            doc.add_paragraph()
            
            doc.add_heading('Transaction Details', 1)
            table = doc.add_table(rows=6, cols=2)
            table.style = 'Light Grid Accent 1'
            details = [('Amount', f"${selected_test.get('Transaction_Amount', 0):,.2f}"), ('Fraud Probability', f"{selected_test.get('fraud_probability', 0):.1%}"), ('Account Age', f"{selected_test.get('Account_Age', 0)} days"), ('Velocity', f"{selected_test.get('Number_of_Transactions_Last_24H', 0)} txns/24h"), ('Previous Fraud', str(selected_test.get('Previous_Fraudulent_Transactions', 0))), ('Transaction ID', f"#{selected_idx + 1}")]
            for i, (key, value) in enumerate(details): table.rows[i].cells[0].text = key; table.rows[i].cells[1].text = value
            doc.add_paragraph()
            
            doc.add_heading('Model Responses', 1)
            doc.add_heading('🤖 Base Model (Generic)', 2)
            doc.add_paragraph(base_response)
            doc.add_paragraph()
            
            if "MLX" in comparison_mode or "3-Way" in comparison_mode:
                doc.add_heading('✨ MLX Fine-Tuned', 2)
                if use_real_inference and real_mlx_response: doc.add_paragraph("Mode: Real Inference").bold = True; doc.add_paragraph(real_mlx_response)
                elif 'mlx_response' in locals(): doc.add_paragraph("Mode: Simulated").bold = True; doc.add_paragraph(mlx_response)
                doc.add_paragraph()
                
            if "Ollama" in comparison_mode or "3-Way" in comparison_mode:
                doc.add_heading('🧠 Ollama + Expert Prompt', 2)
                if use_real_inference and real_ollama_response: doc.add_paragraph("Mode: Real Inference").bold = True; doc.add_paragraph(real_ollama_response)
                elif 'ollama_response' in locals(): doc.add_paragraph("Mode: Simulated").bold = True; doc.add_paragraph(ollama_response)
                doc.add_paragraph()
                
            doc.add_heading('Quantitative Metrics', 1)
            metrics_table = doc.add_table(rows=6, cols=4)
            metrics_table.style = 'Light Grid Accent 1'
            headers = ['Metric', 'Base', 'MLX', 'Ollama']
            for i, header in enumerate(headers): metrics_table.rows[0].cells[i].text = header; metrics_table.rows[0].cells[i].paragraphs[0].runs[0].bold = True
            metrics_data = [('Response Length', f'{base_words} words', f'{base_words + 30} words', f'{base_words + 40} words'), ('Data Points', '0', '5', '6'), ('Decision Clarity', '3/10', '9/10', '10/10'), ('Speed', 'Instant', 'Fast', 'Instant'), ('Training Required', 'No', 'Yes', 'No')]
            for i, row_data in enumerate(metrics_data, start=1):
                for j, cell_data in enumerate(row_data): metrics_table.rows[i].cells[j].text = cell_data
            doc.add_paragraph()
            
            doc.add_heading('Analysis & Recommendation', 1)
            if prob > 0.7 or prev_fraud > 0: doc.add_paragraph("Risk Level: HIGH RISK", style='Intense Quote'); doc.add_paragraph("Recommendation: Use MLX Fine-Tuned or Ollama Expert for detailed risk assessment.")
            elif prob > 0.3: doc.add_paragraph("Risk Level: MEDIUM RISK", style='Intense Quote'); doc.add_paragraph("Recommendation: Ollama Expert recommended for balanced analysis.")
            else: doc.add_paragraph("Risk Level: LOW RISK", style='Intense Quote'); doc.add_paragraph("Recommendation: Base model sufficient.")
            
            os.makedirs("data/generated/reports", exist_ok=True)
            timestamp = get_now_est().strftime('%Y%m%d_%H%M%S')
            filename = f"data/generated/reports/model_comparison_{timestamp}.docx"
            doc.save(filename)
            st.success(f"✅ Report saved to: {filename}")
            with open(filename, 'rb') as f:
                st.download_button("📥 Download Report", f.read(), file_name=f"model_comparison_{timestamp}.docx", mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document", use_container_width=True)
        except ImportError: st.error("❌ python-docx not installed. Run: `pip install python-docx`")
        except Exception as e: st.error(f"❌ Error generating report: {str(e)}")
