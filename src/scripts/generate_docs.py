import os
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

def add_heading(doc, text, level):
    heading = doc.add_heading(text, level=level)
    heading.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

def add_paragraph(doc, text, bold=False, italic=False):
    p = doc.add_paragraph()
    run = p.add_run(text)
    if bold: run.bold = True
    if italic: run.italic = True
    return p

def generate_documentation():
    doc = Document()

    # ---------------------------------------------------------
    # COVER PAGE
    # ---------------------------------------------------------
    doc.add_heading('Sentinel: Fraud Detection AI Workbench', 0).alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    doc.add_paragraph('\n\n\n')
    doc.add_paragraph('Comprehensive System Documentation', style='Subtitle').alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    doc.add_paragraph('\n\n\n')
    doc.add_paragraph('Architecture, Integrated Design, and User Manual', style='Normal').alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    doc.add_page_break()

    # ---------------------------------------------------------
    # TABLE OF CONTENTS (Placeholder, MS Word auto-generates this)
    # ---------------------------------------------------------
    add_heading(doc, 'Table of Contents', 1)
    doc.add_paragraph('1. Introduction of Applications')
    doc.add_paragraph('2. Integrated Design')
    doc.add_paragraph('3. Security Design')
    doc.add_paragraph('4. UI1 - Home')
    doc.add_paragraph('5. UI2 - Data Hub')
    doc.add_paragraph('6. UI3 - SQL RAG')
    doc.add_paragraph('7. UI4 - Graph RAG')
    doc.add_paragraph('8. UI5 - Multimodal RAG')
    doc.add_paragraph('9. UI6 - Trends')
    doc.add_paragraph('10. UI7 - ML Workflow')
    doc.add_paragraph('11. UI8 - LLM Fine Tuning')
    doc.add_paragraph('12. UI9 - API Hub')
    doc.add_paragraph('13. UI10 - Admin Console')
    doc.add_paragraph('14. Sidebar - Show Demo')
    doc.add_paragraph('15. Admin Console - Load Demo Data')
    doc.add_paragraph('16. Dependency References')
    doc.add_paragraph('17. Glossary')
    doc.add_page_break()

    # ---------------------------------------------------------
    # 1. INTRODUCTION OF APPLICATIONS
    # ---------------------------------------------------------
    add_heading(doc, '1. Introduction of Applications', 1)
    doc.add_paragraph(
        "Sentinel Fraud AI Workbench Pro is a comprehensive, enterprise-grade platform engineered to detect, investigate, and mitigate financial fraud. "
        "It unifies Generative AI and traditional Machine Learning within a single, cohesive interface. The platform serves as a central hub for "
        "investigators, data scientists, and administrators to uncover hidden fraud rings, analyze massive datasets, and fine-tune models safely."
    )
    doc.add_paragraph(
        "At its core, Sentinel leverages a Multi-Agent architecture powered by LangGraph. It provides three primary modalities of Retrieval-Augmented Generation (RAG): "
        "SQL RAG for structured transactional data, Graph RAG for interconnected entity analysis (e.g., tracking money laundering across accounts), and "
        "Multimodal RAG for unstructured evidence like images and audio."
    )
    doc.add_paragraph(
        "By integrating these capabilities with real-time Machine Learning pipelines (Random Forest, XGBoost) and persistent graph databases, "
        "Sentinel allows users to instantly transition from asking free-form natural language questions to deploying highly predictive fraud models."
    )

    # ---------------------------------------------------------
    # 2. INTEGRATED DESIGN
    # ---------------------------------------------------------
    add_heading(doc, '2. Integrated Design', 1)
    
    add_heading(doc, '2.1 System Architecture & Data Flow', 2)
    doc.add_paragraph(
        "The architecture adheres to a strict decoupled pattern: a reactive Streamlit frontend communicates securely via JWT-authenticated REST APIs to a high-concurrency "
        "FastAPI backend. The backend acts as the orchestrator, dynamically routing analytical requests to the LangGraph Multi-Agent system, the local ML pipeline, or directly to the data stores."
    )
    
    # Workflow diagram
    diagram_text = (
        "[ User Interface (Streamlit) ]\n"
        "         │ (HTTPS / JWT / API REST)\n"
        "         ▼\n"
        "[ API Gateway (FastAPI) ] ──▶ [ Admin / Auth / RBAC ]\n"
        "         │\n"
        "         ├──▶ [ Multi-Agent System (LangGraph) ] ──▶ [ Local LLMs (Ollama) ]\n"
        "         │        ├── SQL RAG Agent \n"
        "         │        ├── Graph RAG Agent\n"
        "         │        └── Multimodal Agent\n"
        "         │\n"
        "         └──▶ [ ML Pipeline (scikit-learn) ]\n"
        "                  ├── Auto-Training Pipeline\n"
        "                  └── Real-time Feedback Scoring\n"
        "         │\n"
        "         ▼\n"
        "[ Persistent Data Stores ]\n"
        "  - SQLite (Raw structured data mappings)\n"
        "  - Neo4j (Entity graphs & taxonomic relations)\n"
        "  - FAISS (Vector embeddings for KB scaling)\n"
    )
    p = doc.add_paragraph(diagram_text)
    p.runs[0].font.name = 'Courier New'
    p.runs[0].font.size = Pt(9)

    add_heading(doc, '2.2 UI Design Strategy', 2)
    doc.add_paragraph(
        "The UI architecture is entirely modularized under the `src/views/` directory. Each major feature (e.g., Data Hub, ML Workflow) is encapsulated "
        "in its own package, rendering state-aware tabs dynamically based on user context. The global application employs a 'Crimson Dark Mode' theme, utilizing conditional CSS injections "
        "to highlight critical alerts and interactive visual elements. Session State is rigorously managed to ensure that contextual data (such as a selected AI model, filter parameters, or a specific uploaded dataset) "
        "persists seamlessly as the user navigates fluidly between analytical tools."
    )

    # ---------------------------------------------------------
    # 3. SECURITY DESIGN
    # ---------------------------------------------------------
    add_heading(doc, '3. Security Design', 1)
    doc.add_paragraph(
        "Security operates on a zero-trust model heavily focused on Role-Based Access Control (RBAC) and strict API perimeter defense."
    )
    
    add_heading(doc, '3.1 Login Implementations & RBAC', 2)
    doc.add_paragraph(
        "Authentication uses standard OAuth2 architecture with JSON Web Tokens (JWT). Upon providing valid credentials, the backend (FastAPI) securely verifies password hashes utilizing BCrypt. "
        "It then issues a short-lived access token containing role scopes. The frontend caches this token locally and strictly attaches it to the header of every outbound API request."
    )
    doc.add_paragraph(
        "RBAC is enforced dual-layer. On the frontend, the UI constructs the navigation sidebar by querying the `/admin/permissions` API endpoint, hiding pages the user cannot access. "
        "On the backend, routes enforce access internally using parameterized dependency injection (`RequiresRole`), terminating unauthorized requests before any business logic executes."
    )

    add_heading(doc, '3.2 Data Security (SQL Anti-Injection)', 2)
    doc.add_paragraph(
        "The SQL RAG module employs an explicit Abstract Syntax Tree (AST) validator utilizing the `sqlglot` dependency. It inherently rejects LLM-hallucinated queries attempting "
        "destructive DROP, DELETE, UPDATE, INSERT, or ALTER commands. It further restricts queries to explicitly approved reporting tables, fully neutralizing prompt injection vulnerabilities."
    )

    # Helper for UI Chapters
    # Helper for UI Chapters
    def add_ui_section(doc, chapter_num, title, ui_desc, manual_desc, workflow_desc=None, functionality_desc=None):
        add_heading(doc, f'{chapter_num}. {title}', 1)
        add_heading(doc, f'{chapter_num}.1 Functionality Design', 2)
        if functionality_desc:
            doc.add_paragraph(functionality_desc)
        else:
            doc.add_paragraph("This module integrates seamlessly with the Sentinel backend core, utilizing dynamic Session State propagation to maintain contextual continuity across analytical tabs.")
        add_heading(doc, f'{chapter_num}.2 UI Design', 2)
        doc.add_paragraph(ui_desc)
        add_heading(doc, f'{chapter_num}.3 User Manual', 2)
        doc.add_paragraph(manual_desc)
        if workflow_desc:
            add_heading(doc, f'{chapter_num}.4 Logic Workflow', 2)
            p = doc.add_paragraph(workflow_desc)
            p.runs[0].font.name = 'Courier New'
            p.runs[0].font.size = Pt(9)

    add_ui_section(
        doc, '4', 'UI1 - Home',
        "The Home dashboard utilizes a responsive CSS grid layout mapping to three core investigative pillars: Intelligence, Visualization, and AI Core. It displays dynamic real-time system metrics (like active model count and Neo4j node connections) alongside action buttons governed by the authenticated user's RBAC role.",
        "Step 1: Review the system health on the central dashboard immediately upon login.\nStep 2: If the system returns zero metrics, click the 'Quick Install Demo Data' fast-action button to securely populate the SQLite databases and Neo4j graph.\nStep 3: Click any primary module tile to bypass standard sidebar navigation and jump directly into an investigation.",
        "► [ Authenticate User JWT ] \n  └──► [ Fetch System Metrics API ] \n         └──► [ Authorize Role Dashboard UI ] \n                └──► [ Render Action Cards ]"
    )

    add_ui_section(
        doc, '5', 'UI2 - Data Hub',
        "The Data Hub presents a tab-driven interface (Local Uploads, External Sources, Database Management) utilizing highly responsive file-uploader drop-zones and paginated dataframes. Crucially, it embeds a real-time progress bar array that visually indicates concurrent NLP embedding and graph ingest operations.",
        "Step 1: Open the 'Local Uploads' tab and drag a structural CSV file into the drop-zone.\nStep 2: Once uploaded, the backend auto-determines schema types. Click 'Rebuild Knowledge Base' to slice text content and embed it into FAISS.\nStep 3: Map the CSV columns to Node/Edge identifiers to trigger a Neo4j ingestion pipeline.",
        "► [ Upload User CSV ] \n  └──► [ Parse via Pandas ] \n         ├──► [ Auto-Gen SQLite Table Schema ] \n         ├──► [ Vectorize Text via NLP to FAISS ] \n         └──► [ Map Structural Nodes into Neo4j Graph ]"
    )

    add_ui_section(
        doc, '6', 'UI3 - SQL RAG',
        "Constructed around a specialized split-screen IDE layout: the right-hand sidebar is rigidly fixed displaying dynamic active schema dictionaries and auto-generated sample questions. The primary left pane provides an endless-scroll conversational interface supporting dynamic rendering of Markdown, Pandas DataFrames, and Syntax-highlighted SQL blocks.",
        "Step 1: Select an active target database using the header dropdown selector.\nStep 2: Type a natural language investigation metric (e.g., 'Return all accounts initiating over 5 transfers exceeding $10k in 48 hours').\nStep 3: Review the AI's generated SQL, validate the query logic, and download the returned dataframe as an operational CSV.",
        "► [ Natural Language Prompt ] \n  └──► [ Inject DB Schema Context ] \n         └──► [ LLM Generates SQL ] \n                └──► [ AST Security Validation ] \n                       └──► [ SQLite Execution ] \n                              └──► [ Render Pandas DataFrame ]"
    )

    add_ui_section(
        doc, '7', 'UI4 - Graph RAG',
        "Mechanically mirrors the SQL RAG split-screen but substitutes relational schemas for intuitive Node/Edge taxonomy visualization. It integrates a physics-simulated PyVis network Javascript element directly via an interactive `st.components.v1.html` rendering block, situated immediately below the prompt intake.",
        "Step 1: Choose an active Graph taxonomy matrix.\nStep 2: Query the agent for overlapping attributes (e.g., 'Visualize the sub-graph involving IP Address 192.168.1.5').\nStep 3: The interface retrieves the entity map. Manually drag nodes to rearrange the cluster layout, hover for deep property inspection, and review the AI-narrated summary of the identified fraud syndicate.",
        "► [ User Natural Prompt ] \n  └──► [ Fetch Graph Taxonomy Map ] \n         └──► [ LLM Translates to Cypher ] \n                └──► [ Execute via Neo4j Driver ] \n                       └──► [ Generate Physics Graph Elements ] \n                              └──► [ Render HTML iframe PyVis ]"
    )

    add_ui_section(
        doc, '8', 'UI5 - Multimodal RAG',
        "Segmented cleanly via inner tabs dividing logical intake flows: Audio, Image, and Document ingestion. The interface incorporates built-in HTML5 media players enabling contextual review, alongside real-time transcription loading spinners.",
        "Step 1: Upload a piece of localized evidence, such as a scanned invoice image or a suspicious WAV audio recording.\nStep 2: Once uploaded, wait for the OCR or Whisper transcription block to finalize.\nStep 3: Interact specifically with that media via the chat element, instructing the AI to 'Compare the extracted routing number to standard banking indices'.",
        "► [ Receive Binary Media Payload ] \n  └──► [ Classify Media Modality ] \n         ├──► (If Image) [ Extract Layout Extents via OCR ] \n         └──► (If Audio) [ Transcribe via Whisper Node ] \n                └──► [ Pipe Output to Multimodal LLM Agent ] \n                       └──► [ Output Chat Response ]"
    )

    add_ui_section(
        doc, '9', 'UI6 - Trends and Insights',
        "Serves as the primary BI reporting dashboard. Leverages visually reactive Plotly Express charts scaling automatically to container width. Employs advanced cascading drop-downs for dimensional X/Y axis parameterization, dynamic color grouping, and visual-type toggles (scatter, bar, histogram).",
        "Step 1: Select an ingested target dataset.\nStep 2: Assign mathematical parameters to the X/Y axes to identify macro-trends visually (e.g., fraudulent volume plotted across transaction hours).\nStep 3: Click 'Generate AI Insights'. The platform passes the statistical density bounds to the LLM to output a human-friendly narrative explaining the chart.",
        "► [ Select Target Dataframe ] \n  └──► [ Assign Interactive Chart Dimensionality ] \n         └──► [ Generate Plotly Canvas ] \n                └──► (Optional Trigger) \n                       └──► [ Serialize Stats Payload ] \n                              └──► [ Fetch LLM Chart Narrative Analysis ]"
    )

    add_ui_section(
        doc, '10', 'UI7 - ML Workflow',
        "The interface spans multiple horizontal tabs (Build, Train, Tune, Score, Deploy, Monitor) to gracefully orchestrate the end-to-end Machine Learning lifecycle without overwhelming the analyst. Parameters and hyperparameter tuning fields are housed in expandable accordions. Outputs include visually striking KPIs (Accuracy, F1-Score) and interactive matplotlib/plotly charts for SHAP dependency curves and Confusion Matrices.",
        "Step 1: In the 'Build Pipeline' tab, select a dataset, assign target variables, and configure automated feature engineering (imputation, categorical encoding, SMOTE for imbalance).\nStep 2: In the 'Train Model' tab, choose an algorithm (e.g., Random Forest, XGBoost), configure the Train/Test split, and fit the model to view KPIs and SHAP feature importance.\nStep 3: Move to 'Tune Model' for GridSearchCV hyperparameter tuning.\nStep 4: Use 'Batch Scoring' to apply your serialized model to new datasets, and 'Monitor' to observe data drift and performance decay over time.",
        "► [ Data Prep & Feature Eng (Build) ] \n  └──► [ Split Data & SMOTE Balance ] \n         └──► [ Train Classifier (Train/Tune) ] \n                └──► [ SHAP Interpretability Computations ] \n                       └──► [ Serialize Model (.pkl) & Deploy ] \n                              └──► [ Batch Scoring Execution ]",
        "Functionally designed as an automated, multi-tiered pipeline leveraging scikit-learn and imbalanced-learn. The system decomposes the complexity of traditional data science workflows into sequential state-managed modules. It actively mitigates fraud datasets' notorious class imbalance using SMOTE (Synthetic Minority Over-sampling Technique) before executing model training. Furthermore, game-theoretic SHAP (SHapley Additive exPlanations) values are automatically computed to provide strict regulatory transparency into exactly why the algorithmic decision tree flagged a specific transaction."
    )

    add_ui_section(
        doc, '11', 'UI8 - LLM Fine Tuning',
        "A highly structured 'Command Center' consisting of four distinct tabs: Collect, Review, Train, and Test. The UI balances technical command configurations (LoRA rank, Alpha, Batch limits) with intuitive data viewing capabilities. The 'Test' tab features a dual-chat split-pane, rendering identical subsequent prompts to both a 'Base Model' and a 'Fine-Tuned Model' concurrently for direct A/B visual comparison.",
        "Step 1: Use the 'Collect Dataset' tab to securely aggregate Chat/RAG logs, or synthetically generate instructional JSONL datasets using an LLM teacher.\nStep 2: In 'Review', curate the rows to ensure high-quality question/answer pairs and delete anomalous examples.\nStep 3: In the 'Train' tab, select a Base Model (e.g., LLaMA), configure LoRA parameters, and execute the localized MLX training cycle.\nStep 4: Launch the 'Dual-Chat Test' to interrogate both models simultaneously, verifying that the new heavily-weighted fraud knowledge adheres properly.",
        "► [ Synthesize/Extract JSONL ] \n  └──► [ Curate & Review Schema ] \n         └──► [ Initiate MLX LoRA FineTuner Module ] \n                └──► [ Output Custom Quantized Adapters ] \n                       └──► [ Mount Dual Inference Endpoints ] \n                              └──► [ Concurrent A/B Model Querying ]",
        "Designed to bring localized, secure Large Language Model optimization to the analyst without requiring cloud execution (preventing PII/data leakage). It utilizes Low-Rank Adaptation (LoRA) via Apple's MLX (or PyTorch) framework to modify model adapters in a fraction of traditional compilation time. The functional architecture isolates the dataset generation (Extracting historical investigator prompts or using teacher-LLM synthetics) from the multi-threaded inference testbed, ensuring empirical validation of model hallucination reduction."
    )

    add_ui_section(
        doc, '12', 'UI9 - API Interaction Hub',
        "A highly formalized developer portal aesthetic featuring terminal-dark code block structures with rapid 'Copy' clipboard logic. It seamlessly auto-embeds the native FastAPI Swagger documentation sandbox into an HTML frame.",
        "Step 1: Administer the generation of a long-lived API execution token.\nStep 2: Review the dynamically updating parameter requirements.\nStep 3: Utilize the generated cURL/Python integration scripts to tie external banking infrastructure directly into Sentinel's scoring pipeline, pushing live transactions for millisecond ML evaluation.",
        "► [ Generate Secure Request Token ] \n  └──► [ Store Cryptographic Hash Server-Side ] \n         └──► [ Retrieve Target Endpoint Schemas ] \n                └──► [ Render Integration Code Syntax ] \n                       └──► [ Receive External Scoring JSON Post ]"
    )

    add_ui_section(
        doc, '13', 'UI10 - Admin Console',
        "An operational control matrix exclusively protected by strict JWT RBAC scoping. Employs highly organized `st.data_editor` elements for bulk configuration mapping. All changes exhibit immediate hot-reloading behaviors across universal active sessions.",
        "Step 1: Only accessible by authenticated administrators.\nStep 2: Provide specific OpenAI, DeepSeek, or Cloud API tokens; these map dynamically and overwrite `.env` states instantly.\nStep 3: Provision new analyst user accounts, enforce mandatory password resets, and adjust viewing permissions per UI component via toggles.",
        "► [ Validate Admin Scope Requirements ] \n  └──► [ Load Modular UI Tabs (Keys, Roles, SMTP) ] \n         └──► [ Detect User Grid Modifications ] \n                └──► [ Patch Backend Databases ] \n                       └──► [ Force System Hot-Reload Event ]"
    )

    # ---------------------------------------------------------
    # 14. SIDEBAR - SHOW DEMO
    # ---------------------------------------------------------
    add_heading(doc, '14. Sidebar - Show Demo', 1)
    
    add_heading(doc, '14.1 Description', 2)
    doc.add_paragraph("A dedicated function located in the sidebar designed to present an interactive, pre-populated demonstration of the platform's capabilities without requiring user data upload.")
    
    add_heading(doc, '14.2 Steps', 2)
    doc.add_paragraph("Step 1: Locate the 'Show Demo' toggle or button in the main application sidebar.\n"
                      "Step 2: Activate the demo mode to generate and load synthetic demonstration data into the application.\n"
                      "Step 3: Interact with the populated charts, tables, and graphs to explore the platform's features safely.")
                      
    add_heading(doc, '14.3 Impacts to the System', 2)
    doc.add_paragraph("When activated, the Show Demo function bypasses live data fetching and database queries, injecting static synthetic payloads directly into the Session State. The system operates entirely in-memory for the demo session, temporarily suspending external API calls and persistent writes, ensuring that no experimental changes corrupt the underlying Neo4j or SQLite production data stores.")

    # ---------------------------------------------------------
    # 15. ADMIN CONSOLE - LOAD DEMO DATA
    # ---------------------------------------------------------
    add_heading(doc, '15. Admin Console - Load Demo Data', 1)
    
    add_heading(doc, '15.1 Description', 2)
    doc.add_paragraph("A privileged administrative action designed to bootstrap the system by seeding the underlying databases with a comprehensive suite of interconnected, synthetic fraud data. This allows for immediate, full-scale testing of the platform's analytical capabilities without requiring manual data ingestion.")
    
    add_heading(doc, '15.2 Steps and Location', 2)
    doc.add_paragraph("Step 1: Authenticate as a user with 'Administrator' privileges.\n"
                      "Step 2: Navigate to the 'Admin Console' via the main sidebar.\n"
                      "Step 3: Locate the 'Load Demo Data' button within the Database Management or Provisioning section of the console.\n"
                      "Step 4: Click the button and confirm the action to initiate the initialization pipeline.")
                      
    add_heading(doc, '15.3 Impacts to the System', 2)
    doc.add_paragraph("When executed, this function triggers a multi-stage background pipeline that profoundly alters the system state:\n"
                      "• Relational Database (SQLite/PostgreSQL): Injects thousands of synthetic transaction records, user profiles, and account details representing typical and fraudulent activity patterns.\n"
                      "• Graph Database (Neo4j): Constructs a complex web of interconnected nodes (e.g., Users, IP Addresses, Devices, Bank Accounts) and semantic edges (e.g., 'TRANSFERRED_TO', 'LOGGED_IN_FROM'), explicitly modeling known fraud syndicates for Graph RAG analysis.\n"
                      "• Vector Store (FAISS): Embeds pre-generated textual scenarios, simulated case notes, and policy documents, populating the semantic index to enable context-aware similarity searches and Document RAG functionalities.\n"
                      "Warning: This action writes directly to persistent storage and may overwrite or conflict with existing sandbox data.")

    # ---------------------------------------------------------
    # 16. DEPENDENCY REFERENCES
    # ---------------------------------------------------------
    add_heading(doc, '16. Dependency References', 1)
    
    deps = {
        'FastAPI': "Utilized as the core asynchronous backend API router. Chosen for its automatic OpenAPI integration, high throughput, and strict Pydantic data validation model.",
        'Streamlit': "The primary frontend reactive visualization framework. Chosen to accelerate ML tool delivery by eliminating complex React/Vue boilerplate while maintaining interactive graphs.",
        'SQLAlchemy': "An Object Relational Mapper (ORM). Facilitates database-agnostic modeling, ensuring Sentinel can map securely across simple flat SQLite files or enterprise PostgreSQL clusters.",
        'Neo4j (urllib3/neo4j-driver)': "A dedicated graph database. Integral for calculating multi-hop relationships (fraud rings) that deeply struggle in standard row-based SQL structures.",
        'FAISS': "Facebook AI Similarity Search. A highly optimized C++ vector library wrapped in Python enabling instantaneous semantic similarity retrievals across billions of embedded document chunks.",
        'SQLGlot': "An incredibly efficient SQL parser. Necessary for parsing dangerous LLM-generated strings into AST representations to programmatically guarantee the absence of destructive DROP/ALTER statements before engine execution.",
        'PyVis / NetworkX': "Python-native graph visualization suites. Translates raw underlying neo4j edge coordinates into human-readable, physics-simulated node clusters seamlessly overlaid in Streamlit.",
        'BCrypt': "A computationally heavy cryptographic hashing algorithm explicitly chosen to counteract brute-force hardware attacks when storing registered investigator passwords.",
        'Ollama': "An underlying local inference engine wrapper. Enables Sentinel to rapidly pull, mount, and execute quantized open-weight models locally safely entirely off-network."
    }

    counter = 1
    for k, v in deps.items():
        add_heading(doc, f'16.{counter} {k}', 2)
        doc.add_paragraph(v)
        counter += 1

    # ---------------------------------------------------------
    # 17. GLOSSARY
    # ---------------------------------------------------------
    add_heading(doc, '17. Glossary', 1)
    
    glossary = {
        'RAG (Retrieval-Augmented Generation)': 'A methodology where an LLM is paired with an external knowledge retrieval mechanism (like Vector DBs) to ground its responses in factual, enterprise-specific context, eliminating hallucination.',
        'Knowledge Graph': 'A structured taxonomy of concepts and relations (Nodes and Edges). Essential for depicting complex fraud rings where individuals, IP addresses, and bank accounts overlap.',
        'LLM (Large Language Model)': 'A massive neural network trained on vast quantities of text. Capable of translation, summarization, logical bridging, and code compilation (e.g., text-to-SQL logic).',
        'Fine-Tuning (LoRA)': 'Low-Rank Adaptation. A computationally efficient mathematical shortcut to retrain a massive LLM on specific behavior without altering its fundamental parameter structure.',
        'RBAC (Role-Based Access Control)': 'A security paradigm where access logic is denied inherently unless explicitly permitted via a matrix mapped to the authenticated user\'s registered job function.',
        'JWT (JSON Web Token)': 'A cryptographically signed, compact URL-safe execution credential passed between the client and server to verify session integrity asynchronously.',
        'Vector Embeddings': 'Array structures representing the profound semantic meaning of text/images as mathematical distances, allowing search algorithms to identify context similarities regardless of explicit keyword matches.',
        'SHAP (SHapley Additive exPlanations)': 'A game-theoretic ML model evaluation technique that quantifies and visually explains the exact logic path an algorithmic decision tree utilized to calculate a fraud-score output.'
    }

    for term, definition in glossary.items():
        add_paragraph(doc, f"{term}:", bold=True)
        doc.add_paragraph(definition)

    # Save logic
    output_path = os.path.join(os.path.dirname(os.path.dirname(os.path.dirname(os.path.abspath(__file__)))), "Sentinel_Platform_Documentation.docx")
    doc.save(output_path)
    print(f"Documentation generated successfully at: {output_path}")

if __name__ == "__main__":
    generate_documentation()
